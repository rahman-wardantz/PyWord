import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
from tkinter import font
import os

class PyWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PyWord - MS Word-like App")
        self.filename = None
        self.text_modified = False
        self.text = tk.Text(root, wrap='word', undo=True)
        self.text.pack(expand=1, fill='both')
        self.create_menu()
        self.create_toolbar()
        self.create_statusbar()
        self.text.bind('<KeyRelease>', self.update_statusbar)
        self.text.bind('<ButtonRelease>', self.update_statusbar)
        self.text.bind('<<Modified>>', self.on_modified)
        self.root.protocol("WM_DELETE_WINDOW", self.on_exit)
        self.bold_font = font.Font(self.text, self.text.cget("font"))
        self.bold_font.configure(weight="bold")
        self.italic_font = font.Font(self.text, self.text.cget("font"))
        self.italic_font.configure(slant="italic")
        self.underline_font = font.Font(self.text, self.text.cget("font"))
        self.underline_font.configure(underline=1)
        self.text.tag_configure("bold", font=self.bold_font)
        self.text.tag_configure("italic", font=self.italic_font)
        self.text.tag_configure("underline", font=self.underline_font)
        self.root.bind('<Control-b>', lambda e: self.make_bold())
        self.root.bind('<Control-i>', lambda e: self.make_italic())
        self.root.bind('<Control-u>', lambda e: self.make_underline())
        self.root.bind('<Control-s>', lambda e: self.save_docx())
        self.root.bind('<Control-f>', lambda e: self.find_text())
        self.last_search = None
        self.last_search_idx = None
        self.text.tag_configure('search_highlight', background='yellow')

    def create_menu(self):
        menubar = tk.Menu(self.root)
        filemenu = tk.Menu(menubar, tearoff=0)
        filemenu.add_command(label="New", command=self.new_doc)
        filemenu.add_command(label="Open", command=self.open_docx)
        filemenu.add_command(label="Save", command=self.save_docx)
        filemenu.add_command(label="Save As", command=self.saveas_docx)
        filemenu.add_separator()
        filemenu.add_command(label="Exit", command=self.on_exit)
        menubar.add_cascade(label="File", menu=filemenu)
        self.root.config(menu=menubar)

    def set_font_size(self, size):
        # Optimize: update font only if changed, and keep font family and style
        try:
            size = int(size)
        except Exception:
            return
        current_font = font.Font(font=self.text.cget("font"))
        if current_font['size'] == size:
            return
        # Use full font config to avoid resizing widget
        font_tuple = (current_font.actual('family'), size, current_font.actual('weight'), current_font.actual('slant'), int(current_font.actual('underline')))
        self.text.configure(font=font_tuple)
        self.bold_font.configure(size=size)
        self.italic_font.configure(size=size)
        self.underline_font.configure(size=size)
        self.statusbar.config(font=(current_font.actual('family'), max(8, size-2)))
        # Prevent geometry change by forcing the window size to stay
        self.root.update_idletasks()
        # Only set minsize once, not every time
        if not hasattr(self, '_minsize_set'):
            self.root.minsize(self.root.winfo_width(), self.root.winfo_height())
            self._minsize_set = True
        # Set wrap to word to avoid horizontal scroll when font size increases
        self.text.config(wrap='word')
        # Keep focus on text widget after font change
        self.text.focus_set()

    def create_toolbar(self):
        toolbar = tk.Frame(self.root, bd=1, relief=tk.RAISED, bg='#f0f0f0')
        format_frame = tk.Frame(toolbar, bg='#f0f0f0')
        bold_btn = tk.Button(format_frame, text="Bold", width=7, command=self.make_bold, bg='#e0e0e0')
        bold_btn.pack(side=tk.LEFT, padx=1, pady=2)
        italic_btn = tk.Button(format_frame, text="Italic", width=7, command=self.make_italic, bg='#e0e0e0')
        italic_btn.pack(side=tk.LEFT, padx=1, pady=2)
        underline_btn = tk.Button(format_frame, text="Underline", width=9, command=self.make_underline, bg='#e0e0e0')
        underline_btn.pack(side=tk.LEFT, padx=1, pady=2)
        clear_btn = tk.Button(format_frame, text="Clear Formatting", width=15, command=self.clear_formatting, bg='#e0e0e0')
        clear_btn.pack(side=tk.LEFT, padx=1, pady=2)
        format_frame.pack(side=tk.LEFT, padx=4)
        # Font size dropdown
        font_size_var = tk.IntVar(value=12)
        font_size_menu = tk.OptionMenu(toolbar, font_size_var, *[8, 10, 12, 14, 16, 18, 20, 24, 28, 32], command=self.set_font_size)
        font_size_menu.config(width=4, bg='#e0e0e0')
        font_size_menu.pack(side=tk.LEFT, padx=4, pady=2)
        # Separator
        sep1 = tk.Frame(toolbar, width=2, bg='#cccccc', height=28)
        sep1.pack(side=tk.LEFT, padx=2, pady=2)
        # Find/Replace group
        find_frame = tk.Frame(toolbar, bg='#f0f0f0')
        find_btn = tk.Button(find_frame, text="Find", width=7, command=self.find_text, bg='#e0e0e0')
        find_btn.pack(side=tk.LEFT, padx=1, pady=2)
        findnext_btn = tk.Button(find_frame, text="Find Next", width=10, command=self.find_next, bg='#e0e0e0')
        findnext_btn.pack(side=tk.LEFT, padx=1, pady=2)
        replace_btn = tk.Button(find_frame, text="Replace", width=8, command=self.replace_text, bg='#e0e0e0')
        replace_btn.pack(side=tk.LEFT, padx=1, pady=2)
        replacenext_btn = tk.Button(find_frame, text="Replace Next", width=12, command=self.replace_next, bg='#e0e0e0')
        replacenext_btn.pack(side=tk.LEFT, padx=1, pady=2)
        find_frame.pack(side=tk.LEFT, padx=4)
        toolbar.pack(side=tk.TOP, fill=tk.X, pady=2)
        # Set text widget font and background for a modern look
        self.text.config(bg='#fcfcfc', relief=tk.FLAT, bd=2, insertbackground='#222', wrap='word')

    def create_statusbar(self):
        # Optimize: use same font family as editor, smaller size
        current_font = font.Font(font=self.text.cget("font"))
        self.statusbar = tk.Label(self.root, text="Ln 1, Col 1 | Untitled", anchor='w', bg='#eaeaea', fg='#333', font=(current_font.actual('family'), max(8, current_font['size']-2)))
        self.statusbar.pack(side=tk.BOTTOM, fill=tk.X)

    def update_statusbar(self, event=None):
        row, col = self.text.index(tk.INSERT).split('.')
        fname = os.path.basename(self.filename) if self.filename else "Untitled"
        self.statusbar.config(text=f"Ln {int(row)}, Col {int(col)+1} | {fname}")

    def make_bold(self):
        self.toggle_tag("bold")

    def make_italic(self):
        self.toggle_tag("italic")

    def make_underline(self):
        self.toggle_tag("underline")

    def toggle_tag(self, tag):
        try:
            start, end = self.text.index(tk.SEL_FIRST), self.text.index(tk.SEL_LAST)
            if tag in self.text.tag_names(tk.SEL_FIRST):
                self.text.tag_remove(tag, start, end)
            else:
                self.text.tag_add(tag, start, end)
        except tk.TclError:
            pass  # No selection

    def clear_formatting(self):
        self.text.tag_remove("bold", "1.0", tk.END)
        self.text.tag_remove("italic", "1.0", tk.END)
        self.text.tag_remove("underline", "1.0", tk.END)

    def new_doc(self):
        if self.text_modified and not self.confirm_discard_changes():
            return
        self.text.delete(1.0, tk.END)
        self.filename = None
        self.text_modified = False
        self.update_title()

    def save_docx(self):
        if self.filename:
            if os.path.exists(self.filename):
                if not messagebox.askyesno("Overwrite", f"File {os.path.basename(self.filename)} exists. Overwrite?"):
                    return
            try:
                doc = Document()
                content = self.text.get(1.0, tk.END).strip().split('\n')
                for line in content:
                    doc.add_paragraph(line)
                doc.save(self.filename)
                self.text_modified = False
                self.update_title()
                self.update_statusbar()
                messagebox.showinfo("Saved", f"File saved: {os.path.basename(self.filename)}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save file: {e}")
        else:
            self.saveas_docx()

    def saveas_docx(self):
        filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if filepath:
            self.filename = filepath
            self.save_docx()

    def find_text(self):
        query = simpledialog.askstring("Find", "Enter text to find:")
        if not query:
            return
        self.text.tag_remove('search_highlight', '1.0', tk.END)
        start = '1.0'
        found = False
        while True:
            idx = self.text.search(query, start, stopindex=tk.END, nocase=1)
            if not idx:
                break
            end = f"{idx}+{len(query)}c"
            self.text.tag_add('search_highlight', idx, end)
            start = end
            found = True
        if found:
            self.last_search = query
            self.last_search_idx = '1.0'
            self.find_next()
        else:
            messagebox.showinfo("Find", f'"{query}" not found.')

    def find_next(self):
        if not self.last_search:
            return
        idx = self.text.search(self.last_search, self.text.index(tk.INSERT)+"+1c", stopindex=tk.END, nocase=1)
        if idx:
            end = f"{idx}+{len(self.last_search)}c"
            self.text.tag_remove(tk.SEL, '1.0', tk.END)
            self.text.tag_add(tk.SEL, idx, end)
            self.text.mark_set(tk.INSERT, end)
            self.text.see(idx)
        else:
            messagebox.showinfo("Find", f'No more "{self.last_search}" found.')

    def replace_text(self):
        if not self.last_search:
            query = simpledialog.askstring("Replace", "Find what:")
            if not query:
                return
            self.last_search = query
        else:
            query = self.last_search
        replace_with = simpledialog.askstring("Replace", f"Replace '{query}' with:")
        if replace_with is None:
            return
        count = 0
        idx = '1.0'
        while True:
            idx = self.text.search(query, idx, stopindex=tk.END, nocase=1)
            if not idx:
                break
            end = f"{idx}+{len(query)}c"
            self.text.delete(idx, end)
            self.text.insert(idx, replace_with)
            idx = f"{idx}+{len(replace_with)}c"
            count += 1
        if count:
            messagebox.showinfo("Replace", f"Replaced {count} occurrence(s) of '{query}'.")
        else:
            messagebox.showinfo("Replace", f"'{query}' not found.")
        self.text.tag_remove('search_highlight', '1.0', tk.END)
        self.last_search = None
        self.last_search_idx = None

    def replace_next(self):
        if not self.last_search:
            self.replace_text()
            return
        idx = self.text.search(self.last_search, self.text.index(tk.INSERT)+"+1c", stopindex=tk.END, nocase=1)
        if idx:
            end = f"{idx}+{len(self.last_search)}c"
            self.text.tag_remove(tk.SEL, '1.0', tk.END)
            self.text.tag_add(tk.SEL, idx, end)
            self.text.mark_set(tk.INSERT, end)
            self.text.see(idx)
            replace_with = simpledialog.askstring("Replace Next", f"Replace this '{self.last_search}' with:")
            if replace_with is not None:
                self.text.delete(idx, end)
                self.text.insert(idx, replace_with)
        else:
            messagebox.showinfo("Replace Next", f'No more "{self.last_search}" found.')

    def on_modified(self, event=None):
        self.text_modified = self.text.edit_modified()
        self.update_title()
        self.update_statusbar()
        self.text.tag_remove('search_highlight', '1.0', tk.END)
        self.last_search = None
        self.last_search_idx = None
        self.text.edit_modified(False)

    def confirm_discard_changes(self):
        return messagebox.askyesno("Unsaved Changes", "You have unsaved changes. Discard them?")

    def update_title(self):
        name = self.filename if self.filename else "Untitled"
        mod = "*" if self.text_modified else ""
        self.root.title(f"PyWord - {name}{mod}")

    def on_exit(self):
        if self.text_modified and not self.confirm_discard_changes():
            return
        self.root.destroy()

    def open_docx(self):
        if self.text_modified and not self.confirm_discard_changes():
            return
        filepath = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if filepath:
            try:
                doc = Document(filepath)
                self.text.delete(1.0, tk.END)
                for para in doc.paragraphs:
                    self.text.insert(tk.END, para.text + '\n')
                self.filename = filepath
                self.text_modified = False
                self.update_title()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to open file: {e}")

def main():
    root = tk.Tk()
    app = PyWordApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
